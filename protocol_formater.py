""" To run this code:
   brew install texlive
   brew install pandoc
   and imported python libraries
"""

import datetime
import json

from docx import Document
from docx.shared import Pt, Inches, RGBColor
import pypandoc


template_path = "template.docx"
yandex_gpt_answer = 'unofficial_answer.json'


def create_informal_docx(yandex_gpt_answer: str, output_path: str) -> None:
    with open(yandex_gpt_answer, 'r', encoding='utf-8') as file:
        data = json.load(file)

    # Копирование шаблона неофициального
    with open(template_path, 'rb') as src, open(output_path, 'wb') as dst:
        dst.write(src.read())

    font_name = "Century Gothic"

    doc = Document(output_path)

    title_paragraph = doc.add_paragraph()
    run = title_paragraph.add_run(data["Тема"].upper())

    # Задаем шрифт
    run.font.name = font_name

    # Задаем размер шрифта
    run.font.size = Pt(26)

    run.font.color.rgb = RGBColor(2, 9, 63)
    run.bold = True

    # DATE
    current_date = str(datetime.date.today()).replace('-', '.')
    current_time = datetime.datetime.now().time()
    current_time = current_time.strftime("%H:%M")

    subtitle_paragraph = doc.add_paragraph()

    subtitle_text = subtitle_paragraph.add_run("ПРОТОКОЛ ВСТРЕЧИ")
    subtitle_text.font.name = font_name
    subtitle_text.font.size = Pt(16)
    subtitle_text.font.color.rgb = RGBColor(2, 9, 63)
    subtitle_paragraph.paragraph_format.space_before = Pt(8)
    subtitle_paragraph.paragraph_format.space_after = Pt(20)

    table = doc.add_table(rows=1, cols=2)

    cell1 = table.cell(0, 0)
    cell2 = table.cell(0, 1)
    cell1.text = "ДАТА:"
    cell2.text = current_date

    row = table.add_row()
    cell1 = row.cells[0]
    cell2 = row.cells[1]
    cell1.text = "ВРЕМЯ:"
    cell2.text = current_time

    # Длительность
    row = table.add_row()
    cell1 = row.cells[0]
    cell2 = row.cells[1]
    cell1.text = "ДЛИТЕЛЬНОСТЬ:"
    cell2.text = "00:00"

    # # УЧАСТНИКИ
    # row = table.add_row()
    # cell1 = row.cells[0]
    # cell2 = row.cells[1]
    # cell1.text = "УЧАСТНИКИ:"
    # cell2.text = "ИМЕНА"

    for col in table.columns:
        col.width = Inches(0.5)  # Ширина столбца 2.5 дюйма

    # Устанавливаем высоту строк
    for row in table.rows:
        row.height = Pt(20)  # Высота строки 30 пунктов

    # Форматируем текст в ячейках
    for row in table.rows:
        for cell_id, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(9)
                    if cell_id == 0:
                        run.bold = True

    resume_paragraph = doc.add_paragraph()
    resume_text = resume_paragraph.add_run("РЕЗЮМЕ ОБСУЖДЕНИЙ")
    resume_text.font.name = font_name
    resume_text.font.size = Pt(16)
    resume_text.font.color.rgb = RGBColor(2, 9, 63)
    resume_paragraph.paragraph_format.space_before = Pt(8)
    resume_paragraph.paragraph_format.space_after = Pt(20)

    for idx in range(len(data["Тезисы"])):
        list_paragraph_1 = doc.add_paragraph()
        run_1 = list_paragraph_1.add_run(f"{chr(97 + idx)})       {data["Тезисы"][idx]["Формулировка тезиса"]}")
        list_paragraph_1.paragraph_format.left_indent = 360000  # Отступ для маркера
        list_paragraph_1.paragraph_format.first_line_indent = -360000  # Отступ для текста
        run_1.italic = True
        list_paragraph_1.paragraph_format.space_after = Pt(6)

        list_paragraph_2 = doc.add_paragraph()
        run_2 = list_paragraph_2.add_run(
            f"           Контекст обсуждения: {data["Тезисы"][idx]["Контекст обсуждения"]}")
        list_paragraph_2.paragraph_format.left_indent = 360000  # Отступ для маркера
        list_paragraph_2.paragraph_format.first_line_indent = -360000  # Отступ для текста
        list_paragraph_2.paragraph_format.space_after = Pt(6)
        run_2.italic = True
        run_2.font.color.rgb = RGBColor(68, 125, 194)

        list_paragraph_3 = doc.add_paragraph()
        run_3 = list_paragraph_3.add_run(f"           Время: {data["Тезисы"][idx]["Таймкод"]}")
        list_paragraph_3.paragraph_format.left_indent = 360000  # Отступ для маркера
        list_paragraph_3.paragraph_format.first_line_indent = -360000  # Отступ для текста
        list_paragraph_3.paragraph_format.space_after = Pt(6)
        run_3.italic = True
        run_3.font.color.rgb = RGBColor(68, 125, 194)

    doc.save(output_path)

# pdf пока не работет
def create_informal_pdf(yandex_gpt_answer: str, output_path: str) -> None:
    docx_file = output_path[:-3] + "docx"
    create_informal_docx(yandex_gpt_answer, docx_file)
    output = pypandoc.convert_file(docx_file, 'pdf', outputfile=output_path)


create_informal_pdf(yandex_gpt_answer, "results/output.pdf")

