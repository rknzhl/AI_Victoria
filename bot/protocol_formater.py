""" Переводит ответ LLM модели из .json файла в .docx или .pdf
    с определенными параметрами
"""

import datetime
import json
import subprocess
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


template_path = "template.docx"


def create_informal_docx(yandex_gpt_answer: str, output_path: str, timestamps=False, context=False) -> None:
    """
    Формирование .docx файла с неформальным отчетом о встрече

    yandex_gpt_answer: путь до .json файла с ответом YaGPT
    output_path: путь до файла, где сохранится .docx файл. Требует указание полного пути и имени файла
    timestamps: включать в отчет временные промежутки реплик или нет
    context: включать в отчет вонтекст обждения или нет

    Returns: None
    """
    with open(yandex_gpt_answer, 'r', encoding='utf-8') as file:
        data = json.load(file)

    # Копирование шаблона неофициального
    with open(template_path, 'rb') as src, open(output_path, 'wb') as dst:
        dst.write(src.read())

    font_name = "Century Gothic"

    doc = Document(output_path)

    title_paragraph = doc.add_paragraph()
    run = title_paragraph.add_run(data["Название совещания"].upper())

    # Задаем шрифт
    run.font.name = font_name

    # Задаем размер шрифта
    run.font.size = Pt(26)

    run.font.color.rgb = RGBColor(2, 9, 63)
    run.bold = True

    # DATE
    current_date = str(datetime.date.today()).replace('-', '.')
    current_time = datetime.datetime.now().time()
    current_time = current_time.strftime("%H:%M")

    subtitle_paragraph = doc.add_paragraph()

    subtitle_text = subtitle_paragraph.add_run("ПРОТОКОЛ ВСТРЕЧИ")
    subtitle_text.font.name = font_name
    subtitle_text.font.size = Pt(16)
    subtitle_text.font.color.rgb = RGBColor(2, 9, 63)
    subtitle_paragraph.paragraph_format.space_before = Pt(8)
    subtitle_paragraph.paragraph_format.space_after = Pt(20)

    table = doc.add_table(rows=1, cols=2)

    cell1 = table.cell(0, 0)
    cell2 = table.cell(0, 1)
    cell1.text = "ДАТА:"
    cell2.text = current_date

    row = table.add_row()
    cell1 = row.cells[0]
    cell2 = row.cells[1]
    cell1.text = "ВРЕМЯ:"
    cell2.text = current_time

    # Длительность
    row = table.add_row()
    cell1 = row.cells[0]
    cell2 = row.cells[1]
    cell1.text = "ДЛИТЕЛЬНОСТЬ:"
    cell2.text = "00:00"


    # # УЧАСТНИКИ
    # row = table.add_row()
    # cell1 = row.cells[0]
    # cell2 = row.cells[1]
    # cell1.text = "УЧАСТНИКИ:"
    # cell2.text = "ИМЕНА"

    for col in table.columns:
        col.width = Inches(0.5)  # Ширина столбца 2.5 дюйма

    # Устанавливаем высоту строк
    for row in table.rows:
        row.height = Pt(20)  # Высота строки 30 пунктов

    # Форматируем текст в ячейках
    for row in table.rows:
        for cell_id, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(9)
                    if cell_id == 0:
                        run.bold = True

    goal_paragraph = doc.add_paragraph()
    run = goal_paragraph.add_run("ПОВЕСТКА ДНЯ")

    # Задаем шрифт
    run.font.name = font_name

    # Задаем размер шрифта
    run.font.size = Pt(16)

    run.font.color.rgb = RGBColor(2, 9, 63)

    goal_paragraph = doc.add_paragraph()
    run = goal_paragraph.add_run(data["Цель встречи"])
    run.font.name = font_name
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(2, 9, 63)
    goal_paragraph.paragraph_format.space_before = Pt(8)
    goal_paragraph.paragraph_format.space_after = Pt(20)


    resume_paragraph = doc.add_paragraph()
    resume_text = resume_paragraph.add_run("РЕЗЮМЕ ОБСУЖДЕНИЙ")
    resume_text.font.name = font_name
    resume_text.font.size = Pt(16)
    resume_text.font.color.rgb = RGBColor(2, 9, 63)
    resume_paragraph.paragraph_format.space_before = Pt(8)
    resume_paragraph.paragraph_format.space_after = Pt(20)

    for idx in range(len(data["Тезисы"])):
        list_paragraph_1 = doc.add_paragraph()
        run_1 = list_paragraph_1.add_run(f"{chr(97 + idx)})       {data["Тезисы"][idx]["Формулировка тезиса"]}")
        list_paragraph_1.paragraph_format.left_indent = 360000  # Отступ для маркера
        list_paragraph_1.paragraph_format.first_line_indent = -360000  # Отступ для текста
        run_1.italic = True
        list_paragraph_1.paragraph_format.space_after = Pt(6)

        if context:
            list_paragraph_2 = doc.add_paragraph()
            run_2 = list_paragraph_2.add_run(
                f"           Контекст обсуждения: {data["Тезисы"][idx]["Контекст обсуждения"]}")
            list_paragraph_2.paragraph_format.left_indent = 360000  # Отступ для маркера
            list_paragraph_2.paragraph_format.first_line_indent = -360000  # Отступ для текста
            list_paragraph_2.paragraph_format.space_after = Pt(6)
            run_2.italic = True
            run_2.font.color.rgb = RGBColor(68, 125, 194)

        if timestamps:
            list_paragraph_3 = doc.add_paragraph()
            run_3 = list_paragraph_3.add_run(f"           Время: {data["Тезисы"][idx]["Таймкод"]}")
            list_paragraph_3.paragraph_format.left_indent = 360000  # Отступ для маркера
            list_paragraph_3.paragraph_format.first_line_indent = -360000  # Отступ для текста
            list_paragraph_3.paragraph_format.space_after = Pt(6)
            run_3.italic = True
            run_3.font.color.rgb = RGBColor(68, 125, 194)

    doc.save(output_path)


def create_informal_pdf(yandex_gpt_answer: str, output_path: str, timestamps=False, context=False) -> None:
    """
    Формирование .pdf файла с неформальным отчетом о встрече.

    yandex_gpt_answer: путь до .json файла с ответом YaGPT
    output_path: путь до файла, где сохранится .docx файл. Требует указание полного пути и имени файла
    timestamps: включать в отчет временные промежутки реплик или нет
    context: включать в отчет вонтекст обждения или нет

    Returns: None
    """
    docx_file = output_path[:-3] + "docx"
    create_informal_docx(yandex_gpt_answer, docx_file, timestamps, context)

    shell_script = './convert_docx_to_pdf.sh'
    result = subprocess.run(['bash', shell_script, docx_file], capture_output=True, text=True)



def create_formal_docx(yandex_gpt_answer: str, output_path: str, timestamps=False, context=False, assignment=False) -> None:
    """
        Формирование .docx файла с формальным отчетом о встрече

        yandex_gpt_answer: путь до .json файла с ответом YaGPT
        output_path: путь до файла, где сохранится .docx файл. Требует указание полного пути и имени файла
        timestamps: включать в отчет временные промежутки реплик или нет
        context: включать в отчет вонтекст обждения или нет

        Returns: None
        """
    with open(yandex_gpt_answer, 'r', encoding='utf-8') as file:
        data = json.load(file)

    # Копирование шаблона неофициального
    with open(template_path, 'rb') as src, open(output_path, 'wb') as dst:
        dst.write(src.read())

    font_name = "Times New Roman"

    doc = Document()

    protocol_name = data["Название протокола"].split(" ")
    protocol_word = protocol_name[0].upper()
    protocol_name = " ".join(protocol_name[1:])

    protocol_paragraph = doc.add_paragraph()
    run = protocol_paragraph.add_run(protocol_word)
    run.font.name = font_name
    run.font.size = Pt(14)
    run.bold = True
    protocol_paragraph.paragraph_format.space_before = Pt(8)
    protocol_paragraph.paragraph_format.space_after = Pt(8)
    protocol_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_paragraph = doc.add_paragraph()
    run = title_paragraph.add_run(protocol_name)
    run.font.name = font_name
    run.font.size = Pt(14)
    run.bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_paragraph.paragraph_format.left_indent = Pt(20)  # Отступ слева
    line_paragraph.paragraph_format.right_indent = Pt(20)  # Отступ справа
    run = line_paragraph.add_run("________________________________________________________")
    run.font.name = font_name
    run.font.size = Pt(14)

    city_paragraph = doc.add_paragraph()
    run = city_paragraph.add_run("Москва")
    run.font.name = font_name
    run.font.size = Pt(14)
    run.bold = True
    city_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # DATE
    current_date = str(datetime.date.today()).replace('-', '.')
    current_time = datetime.datetime.now().time()

    date_paragraph = doc.add_paragraph()
    run = date_paragraph.add_run(current_date + "№___________")
    run.font.name = font_name
    run.font.size = Pt(14)
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_paragraph.paragraph_format.left_indent = Pt(20)  # Отступ слева
    line_paragraph.paragraph_format.right_indent = Pt(20)  # Отступ справа
    run = line_paragraph.add_run("________________________________________________________")
    run.font.name = font_name
    run.font.size = Pt(14)

    theme_paragraph = doc.add_paragraph()
    run = theme_paragraph.add_run(data["Тема"])
    run.font.name = font_name
    run.font.size = Pt(14)
    theme_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_paragraph.paragraph_format.left_indent = Pt(20)  # Отступ слева
    line_paragraph.paragraph_format.right_indent = Pt(20)  # Отступ справа
    run = line_paragraph.add_run("________________________________________________________")
    run.font.name = font_name
    run.font.size = Pt(14)

    for idx in range(len(data["Постановления/решения"])):
        list_paragraph_1 = doc.add_paragraph()
        run_1 = list_paragraph_1.add_run(f"{idx + 1}.     {data["Постановления/решения"][idx]["Формулировка"]}")
        list_paragraph_1.paragraph_format.left_indent = 360000  # Отступ для маркера
        list_paragraph_1.paragraph_format.first_line_indent = -360000  # Отступ для текста
        list_paragraph_1.paragraph_format.space_after = Pt(6)
        run_1.font.name = font_name
        run_1.font.size = Pt(14)

        if context:
            list_paragraph_2 = doc.add_paragraph()
            run_2 = list_paragraph_2.add_run(
                f"        Контекст обсуждения: {data["Постановления/решения"][idx]["Контекст обсуждения"]}")
            list_paragraph_2.paragraph_format.left_indent = 360000  # Отступ для маркера
            list_paragraph_2.paragraph_format.first_line_indent = -360000  # Отступ для текста
            list_paragraph_2.paragraph_format.space_after = Pt(6)
            run_2.italic = True
            run_2.font.color.rgb = RGBColor(68, 125, 194)
            run_2.font.name = font_name
            run_2.font.size = Pt(14)

        if timestamps:
            list_paragraph_3 = doc.add_paragraph()
            run_3 = list_paragraph_3.add_run(f"        Время: {data["Постановления/решения"][idx]["Таймкод"]}")
            list_paragraph_3.paragraph_format.left_indent = 360000  # Отступ для маркера
            list_paragraph_3.paragraph_format.first_line_indent = -360000  # Отступ для текста
            list_paragraph_3.paragraph_format.space_after = Pt(6)
            run_3.italic = True
            run_3.font.color.rgb = RGBColor(68, 125, 194)
            run_3.font.name = font_name
            run_3.font.size = Pt(14)

    if assignment:
        assignment_paragraph = doc.add_paragraph()
        run = assignment_paragraph.add_run("ПЕРЕЧЕНЬ ПОРУЧЕНИЙ")
        run.font.name = font_name
        run.font.size = Pt(14)
        run.bold = True
        assignment_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for idx in range(len(data["Перечень поручений"])):
            list_paragraph_1 = doc.add_paragraph()
            run_1 = list_paragraph_1.add_run(f"{idx + 1}.     {data["Перечень поручений"][idx]["Поручение"]}")
            list_paragraph_1.paragraph_format.left_indent = 360000  # Отступ для маркера
            list_paragraph_1.paragraph_format.first_line_indent = -360000  # Отступ для текста
            list_paragraph_1.paragraph_format.space_after = Pt(6)
            run_1.font.name = font_name
            run_1.font.size = Pt(14)


    doc.save(output_path)



def create_formal_pdf(yandex_gpt_answer: str, output_path: str, timestamps=False, context=False, assignment=False) -> None:
    """
    Формирование .pdf файла с формальным отчетом о встрече.

    yandex_gpt_answer: путь до .json файла с ответом YaGPT
    output_path: путь до файла, где сохранится .docx файл. Требует указание полного пути и имени файла
    timestamps: включать в отчет временные промежутки реплик или нет
    context: включать в отчет вонтекст обждения или нет

    Returns: None
    """
    docx_file = output_path[:-3] + "docx"
    create_formal_docx(yandex_gpt_answer, docx_file, timestamps, context, assignment)

    shell_script = './convert_docx_to_pdf.sh'
    result = subprocess.run(['bash', shell_script, docx_file], capture_output=True, text=True)
